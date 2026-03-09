import os
import json
import uuid
import traceback

from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import List, Dict, Optional

# Document Generation
from fpdf import FPDF
from docx import Document

# AI Content Engine
from resume_ai_engine import generate_smart_bullets, generate_smart_summary, _match_role

router = APIRouter()

GENERATED_FILES_DIR = "static/resumes"
os.makedirs(GENERATED_FILES_DIR, exist_ok=True)

class ChatMessage(BaseModel):
    role: str
    content: str

class ImproveResumeRequest(BaseModel):
    uid: str
    messages: List[ChatMessage]

# ─── 10 Professional ATS Resume Templates ───
RESUME_TEMPLATES = [
    {"id": "classic",   "name": "Classic Professional", "accent": (0, 51, 102),   "header_bg": (0, 51, 102)},
    {"id": "modern",    "name": "Modern Teal",          "accent": (0, 128, 128),  "header_bg": (0, 128, 128)},
    {"id": "executive", "name": "Executive Gold",       "accent": (139, 90, 43),  "header_bg": (51, 51, 51)},
    {"id": "minimal",   "name": "Clean Minimal",        "accent": (80, 80, 80),   "header_bg": (60, 60, 60)},
    {"id": "tech",      "name": "Tech Blue",            "accent": (30, 100, 200), "header_bg": (25, 55, 109)},
    {"id": "creative",  "name": "Creative Red",         "accent": (180, 40, 40),  "header_bg": (140, 30, 30)},
    {"id": "academic",  "name": "Academic Green",       "accent": (0, 100, 60),   "header_bg": (0, 80, 50)},
    {"id": "startup",   "name": "Startup Purple",       "accent": (100, 50, 150), "header_bg": (80, 40, 120)},
    {"id": "corporate", "name": "Corporate Navy",       "accent": (0, 40, 85),    "header_bg": (0, 30, 70)},
    {"id": "elegant",   "name": "Elegant Charcoal",     "accent": (50, 50, 50),   "header_bg": (40, 40, 40)},
]

def _safe(text: str) -> str:
    if not isinstance(text, str):
        text = str(text)
    return text.encode('ascii', 'replace').decode('ascii')


# ─── Professional PDF Generator ───

def generate_pdf(data: dict, filename: str, template: dict = None) -> str:
    import random
    if template is None:
        template = random.choice(RESUME_TEMPLATES)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    accent = template["accent"]
    header_bg = template["header_bg"]
    
    # ── Header Block with colored background ──
    pdf.set_fill_color(*header_bg)
    pdf.rect(0, 0, 210, 42, 'F')
    
    # Name (white on dark header)
    pdf.set_y(8)
    pdf.set_font("Helvetica", 'B', 24)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(w=0, h=12, text=_safe(data.get("name", "").upper()), align='C')
    pdf.ln(14)
    
    # Contact row
    pdf.set_font("Helvetica", '', 9)
    pdf.set_text_color(220, 220, 220)
    contact_parts = []
    if data.get("email"):
        contact_parts.append(_safe(data["email"]))
    if data.get("phone"):
        contact_parts.append(_safe(data["phone"]))
    if data.get("linkedin"):
        contact_parts.append(_safe(data["linkedin"]))
    if data.get("github"):
        contact_parts.append(_safe(data["github"]))
    if data.get("portfolio"):
        contact_parts.append(_safe(data["portfolio"]))
    contact_line = "  |  ".join(contact_parts)
    pdf.cell(w=0, h=6, text=contact_line, align='C')
    pdf.ln(14)
    
    # Reset text to black
    pdf.set_text_color(0, 0, 0)
    y_after_header = pdf.get_y()
    
    def _section_header(title):
        pdf.ln(3)
        pdf.set_font("Helvetica", 'B', 11)
        pdf.set_text_color(*accent)
        pdf.cell(w=0, h=7, text=title.upper())
        pdf.ln(8)
        pdf.set_draw_color(*accent)
        pdf.set_line_width(0.6)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(4)
    
    # ── Professional Summary ──
    if data.get("summary"):
        _section_header("Professional Summary")
        pdf.set_font("Helvetica", '', 9)
        pdf.set_text_color(40, 40, 40)
        pdf.multi_cell(w=0, h=4.5, text=_safe(data["summary"]))
    
    # ── Work Experience ──
    if data.get("experience"):
        _section_header("Work Experience")
        for exp in data["experience"]:
            pdf.set_font("Helvetica", 'B', 10)
            pdf.set_text_color(0, 0, 0)
            title = _safe(exp.get("title", ""))
            company = _safe(exp.get("company", ""))
            # Title and company on same line, date right-aligned
            pdf.cell(w=140, h=6, text=f"{title} | {company}")
            pdf.set_font("Helvetica", 'I', 9)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(w=0, h=6, text=_safe(exp.get("date", "")), align='R')
            pdf.ln(7)
            
            pdf.set_font("Helvetica", '', 9)
            pdf.set_text_color(40, 40, 40)
            for bullet in exp.get("bullets", []):
                pdf.cell(w=5, h=4.5, text="")  # indent
                pdf.multi_cell(w=0, h=4.5, text="- " + _safe(bullet))
                pdf.ln(0.5)
            pdf.ln(2)
    
    # ── Education ──
    if data.get("education"):
        _section_header("Education")
        for edu in data["education"]:
            pdf.set_font("Helvetica", 'B', 10)
            pdf.set_text_color(0, 0, 0)
            degree = _safe(edu.get("degree", ""))
            school = _safe(edu.get("school", ""))
            pdf.cell(w=140, h=6, text=f"{degree}")
            pdf.set_font("Helvetica", 'I', 9)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(w=0, h=6, text=_safe(edu.get("date", "")), align='R')
            pdf.ln(6)
            pdf.set_font("Helvetica", '', 9)
            pdf.set_text_color(60, 60, 60)
            pdf.cell(w=0, h=5, text=school)
            pdf.ln(4)
            if edu.get("gpa"):
                pdf.set_font("Helvetica", 'I', 9)
                pdf.cell(w=0, h=5, text=f"GPA: {_safe(edu['gpa'])}")
                pdf.ln(5)
            pdf.ln(1)
    
    # ── Technical Skills ──
    if data.get("skills"):
        _section_header("Technical Skills")
        pdf.set_font("Helvetica", '', 9)
        pdf.set_text_color(40, 40, 40)
        if isinstance(data["skills"], list):
            # Group skills nicely
            skills_text = _safe("  |  ".join(data["skills"]))
        else:
            skills_text = _safe(str(data["skills"]))
        pdf.multi_cell(w=0, h=5, text=skills_text)
    
    # ── Projects ──
    if data.get("projects"):
        _section_header("Projects")
        for proj in data["projects"]:
            pdf.set_font("Helvetica", 'B', 10)
            pdf.set_text_color(0, 0, 0)
            pdf.cell(w=0, h=6, text=_safe(proj.get("name", "")))
            pdf.ln(6)
            pdf.set_font("Helvetica", '', 9)
            pdf.set_text_color(40, 40, 40)
            pdf.multi_cell(w=0, h=4.5, text="- " + _safe(proj.get("description", "")))
            pdf.ln(2)
    
    # ── Certifications ──
    if data.get("certifications"):
        _section_header("Certifications")
        pdf.set_font("Helvetica", '', 9)
        pdf.set_text_color(40, 40, 40)
        for cert in data["certifications"]:
            pdf.cell(w=0, h=5, text="- " + _safe(cert))
            pdf.ln(5)

    filepath = os.path.join(GENERATED_FILES_DIR, f"{filename}.pdf")
    pdf.output(filepath)
    return f"/static/resumes/{filename}.pdf"


def generate_docx(data: dict, filename: str) -> str:
    doc = Document()
    
    doc.add_heading(data.get("name", ""), 0)
    contact_parts = []
    for key in ["email", "phone", "linkedin", "github", "portfolio"]:
        if data.get(key):
            contact_parts.append(data[key])
    doc.add_paragraph(" | ".join(contact_parts))
    
    if data.get("summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(data["summary"])
        
    if data.get("experience"):
        doc.add_heading("Work Experience", level=1)
        for exp in data["experience"]:
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('title', '')} | {exp.get('company', '')}").bold = True
            doc.add_paragraph(exp.get("date", ""), style='Intense Quote')
            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style='List Bullet')

    if data.get("education"):
        doc.add_heading("Education", level=1)
        for edu in data["education"]:
            doc.add_paragraph(f"{edu.get('degree', '')} - {edu.get('school', '')} ({edu.get('date', '')})")
            if edu.get("gpa"):
                doc.add_paragraph(f"GPA: {edu['gpa']}")

    if data.get("skills"):
        doc.add_heading("Technical Skills", level=1)
        if isinstance(data["skills"], list):
            doc.add_paragraph(" | ".join(data["skills"]))
        else:
            doc.add_paragraph(str(data["skills"]))
    
    if data.get("projects"):
        doc.add_heading("Projects", level=1)
        for proj in data["projects"]:
            p = doc.add_paragraph()
            p.add_run(proj.get("name", "")).bold = True
            doc.add_paragraph(proj.get("description", ""), style='List Bullet')
    
    if data.get("certifications"):
        doc.add_heading("Certifications", level=1)
        for cert in data["certifications"]:
            doc.add_paragraph(cert, style='List Bullet')

    filepath = os.path.join(GENERATED_FILES_DIR, f"{filename}.docx")
    doc.save(filepath)
    return f"/static/resumes/{filename}.docx"


# ─── Multi-Step Conversation Engine ───
# Each session stores collected data keyed by uid

_session_data = {}  # uid -> dict of collected resume fields

STEPS = [
    {"key": "name",           "question": "Let's build your professional resume! First, what is your **full name**?"},
    {"key": "email",          "question": "Great! What is your **email address**?"},
    {"key": "phone",          "question": "What is your **phone number**? (with country code, e.g., +91 9876543210)"},
    {"key": "linkedin",       "question": "Do you have a **LinkedIn profile URL**? (type 'skip' if you don't have one)"},
    {"key": "target_role",    "question": "What **role** are you applying for? (e.g., 'Software Engineer', 'Frontend Developer')"},
    {"key": "target_company", "question": "Which **company** are you targeting? (e.g., 'Google', 'Amazon', 'TCS')"},
    {"key": "education",      "question": "Tell me about your **education**. Please provide:\n**Degree, College/University, Year** (e.g., 'B.Tech CSE, VIT University, 2021-2025')\n\nYou can add multiple entries separated by semicolons (;)"},
    {"key": "experience",     "question": "Do you have any **work experience or internships**? If yes, provide:\n**Title, Company, Duration** (e.g., 'Frontend Intern, Infosys, Jun 2024 - Aug 2024')\n\nType 'fresher' if you have no experience. Separate multiple entries with semicolons (;)"},
    {"key": "skills",         "question": "List your **technical skills** separated by commas:\n(e.g., 'Python, React, AWS, SQL, Git, Docker')"},
    {"key": "projects",       "question": "Mention 1-2 **notable projects** with a brief description:\n(e.g., 'E-Commerce App - Built a full-stack shopping platform using React and Node.js')\n\nSeparate multiple projects with semicolons (;). Type 'skip' to skip."},
    {"key": "certifications",  "question": "Any **certifications**? (e.g., 'AWS Cloud Practitioner, Google IT Support')\n\nType 'skip' if none."},
]


def _get_step_index(uid: str) -> int:
    """Get current step for this user session."""
    if uid not in _session_data:
        _session_data[uid] = {"_step": 0}
    return _session_data[uid].get("_step", 0)


def _store_answer(uid: str, key: str, value: str):
    """Store user answer for a step."""
    if uid not in _session_data:
        _session_data[uid] = {"_step": 0}
    _session_data[uid][key] = value.strip()
    _session_data[uid]["_step"] = _session_data[uid].get("_step", 0) + 1


def _parse_education(text: str) -> list:
    """Parse education entries from user text."""
    entries = [e.strip() for e in text.split(";") if e.strip()]
    result = []
    for entry in entries:
        parts = [p.strip() for p in entry.split(",")]
        edu = {
            "degree": parts[0] if len(parts) > 0 else entry,
            "school": parts[1] if len(parts) > 1 else "",
            "date":   parts[2] if len(parts) > 2 else "",
        }
        if len(parts) > 3:
            edu["gpa"] = parts[3]
        result.append(edu)
    return result if result else [{"degree": text, "school": "", "date": ""}]


def _parse_experience(text: str, skills: list, role: str = "", company: str = "") -> list:
    """Parse experience entries and generate AI-driven bullet points."""
    if text.lower().strip() in ["fresher", "no", "none", "skip", "na", "n/a"]:
        bullets = generate_smart_bullets(role or "intern", company or "", skills, count=5)
        return [{
            "title": "Software Development (Academic Projects)",
            "company": "University / Personal Projects",
            "date": "2024 - Present",
            "bullets": bullets
        }]
    
    entries = [e.strip() for e in text.split(";") if e.strip()]
    result = []
    for entry in entries:
        parts = [p.strip() for p in entry.split(",")]
        # Use the specific entry title for bullet generation if it mentions a role
        entry_role = parts[0] if len(parts) > 0 else role
        entry_company = parts[1] if len(parts) > 1 else company
        
        bullets = generate_smart_bullets(
            role=entry_role,
            company=entry_company,
            skills=skills,
            count=5
        )
        
        exp = {
            "title":   parts[0] if len(parts) > 0 else entry,
            "company": parts[1] if len(parts) > 1 else "",
            "date":    parts[2] if len(parts) > 2 else "",
            "bullets": bullets
        }
        result.append(exp)
    return result if result else [{"title": text, "company": "", "date": "", "bullets": generate_smart_bullets(role, company, skills)}]


def _parse_projects(text: str) -> list:
    if text.lower().strip() in ["skip", "no", "none", "na", "n/a"]:
        return []
    entries = [e.strip() for e in text.split(";") if e.strip()]
    result = []
    for entry in entries:
        parts = entry.split(" - ", 1)
        result.append({
            "name": parts[0].strip(),
            "description": parts[1].strip() if len(parts) > 1 else parts[0].strip()
        })
    return result


def _parse_certifications(text: str) -> list:
    if text.lower().strip() in ["skip", "no", "none", "na", "n/a"]:
        return []
    return [c.strip() for c in text.split(",") if c.strip()]


def _build_resume_data(uid: str) -> dict:
    """Build the final resume data from collected session info."""
    sd = _session_data.get(uid, {})
    
    skills = [s.strip() for s in sd.get("skills", "").split(",") if s.strip()]
    role = sd.get("target_role", "Software Developer")
    company = sd.get("target_company", "")
    
    experience_text = sd.get("experience", "fresher")
    exp_years = ""
    if experience_text.lower().strip() not in ["fresher", "no", "none", "skip"]:
        exp_years = "1+"  # Default assumption
    else:
        exp_years = "fresher"
    
    education = _parse_education(sd.get("education", ""))
    experience = _parse_experience(experience_text, skills, role, company)
    projects = _parse_projects(sd.get("projects", "skip"))
    certifications = _parse_certifications(sd.get("certifications", "skip"))
    
    # AI-driven summary generation
    summary = generate_smart_summary(
        name=sd.get("name", ""),
        role=role,
        company=company,
        skills=skills,
        exp_text=experience_text
    )
    
    # Get AI skill suggestions for the role
    role_data = _match_role(role)
    suggested_skills = role_data.get("skills_suggest", [])
    # Merge user skills with suggested ones (user's first)
    all_skills = list(dict.fromkeys(skills + [s for s in suggested_skills if s.lower() not in [sk.lower() for sk in skills]]))
    
    return {
        "name": sd.get("name", ""),
        "email": sd.get("email", ""),
        "phone": sd.get("phone", ""),
        "linkedin": sd.get("linkedin", "") if sd.get("linkedin", "").lower() != "skip" else "",
        "summary": summary,
        "experience": experience,
        "education": education,
        "skills": all_skills,
        "projects": projects,
        "certifications": certifications,
    }



@router.post("/improve-resume")
async def improve_resume(req: ImproveResumeRequest):
    try:
        uid = req.uid or "default"
        
        # Filter system context messages
        user_messages = [m for m in req.messages if m.role == "user"
                         and not m.content.startswith("System Context:")]
        last_msg = user_messages[-1].content.strip().lower() if user_messages else ""
        
        # ── Initialise session data ────────────────────────────────────────────
        if uid not in _session_data:
            _session_data[uid] = {"_step": 0, "_mode": None}
        
        session = _session_data[uid]
        mode = session.get("_mode")
        
        # ── Extract profile context (skills, ATS) from System Context message ──
        for m in req.messages:
            if m.role == "user" and m.content.startswith("System Context:"):
                if "skills are:" in m.content:
                    skills_part = m.content.split("skills are:")[1].split(".")[0].strip()
                    session["_profile_skills"] = skills_part
                if "resume score is" in m.content:
                    score_part = m.content.split("resume score is")[1].split("/")[0].strip()
                    try:
                        session["_ats_score"] = int(score_part)
                    except ValueError:
                        pass
        
        # ── Reset command ──────────────────────────────────────────────────────
        if last_msg in ["restart", "start over", "reset", "new", "menu", "back"]:
            _session_data[uid] = {"_step": 0, "_mode": None}
            session = _session_data[uid]
            mode = None
        
        # ══ STEP 0: First greeting — show Rate or Build options ════════════════
        if len(user_messages) == 0 or (len(user_messages) == 1 and mode is None and session.get("_step", 0) == 0):
            _session_data[uid] = {"_step": 0, "_mode": None,
                                  "_profile_skills": session.get("_profile_skills", ""),
                                  "_ats_score": session.get("_ats_score", 0)}
            return {
                "reply": (
                    "Hello! 👋 I'm your **AI Resume Assistant**.\n\n"
                    "What would you like me to do today?\n\n"
                    "**1️⃣ Rate my resume** — I'll analyze it and give you detailed improvement suggestions.\n"
                    "**2️⃣ Build me a new resume** — I'll collect your information step-by-step and generate a professional ATS-optimized resume.\n\n"
                    "Just reply with **1**, **rate**, **2**, or **build** to get started! 🚀"
                ),
                "pdf_url": None,
                "docx_url": None
            }
        
        # ══ STEP 1: User selects mode ══════════════════════════════════════════
        if mode is None:
            is_rate = any(w in last_msg for w in ["1", "rate", "analyze", "score", "check", "review", "evaluate"])
            is_build = any(w in last_msg for w in ["2", "build", "create", "make", "generate", "new resume"])
            
            if is_rate:
                session["_mode"] = "rate"
                session["_rate_step"] = 0
                return {
                    "reply": (
                        "Great! Let's **rate your resume**. 📊\n\n"
                        "Would you like me to analyze:\n\n"
                        "**A) Your profile resume** — the one you already uploaded in your Profile section\n"
                        "**B) A new resume** — paste or upload a different resume\n\n"
                        "Reply with **A** or **B**."
                    ),
                    "pdf_url": None,
                    "docx_url": None
                }
            elif is_build:
                session["_mode"] = "build"
                session["_step"] = 0
                return {
                    "reply": "Let's build your resume! 📄\n\n" + STEPS[0]["question"],
                    "pdf_url": None,
                    "docx_url": None
                }
            else:
                return {
                    "reply": (
                        "I didn't quite catch that. Please reply with:\n\n"
                        "**1** or **rate** — to analyze your resume\n"
                        "**2** or **build** — to create a new resume"
                    ),
                    "pdf_url": None,
                    "docx_url": None
                }
        
        # ══ RATE MODE ═══════════════════════════════════════════════════════════
        if mode == "rate":
            rate_step = session.get("_rate_step", 0)
            
            # Rate step 0: user picks A or B
            if rate_step == 0:
                is_profile = any(w in last_msg for w in ["a", "profile", "uploaded", "existing", "my resume"])
                is_new = any(w in last_msg for w in ["b", "new", "paste", "another", "upload", "different"])
                
                if is_profile:
                    session["_rate_step"] = 2  # profile path
                    profile_skills = session.get("_profile_skills", "")
                    ats_score = session.get("_ats_score", 0)
                    
                    if not profile_skills and ats_score == 0:
                        return {
                            "reply": (
                                "⚠️ I couldn't find your uploaded profile resume data.\n\n"
                                "Please make sure you have uploaded a resume in your **Profile** page first.\n\n"
                                "Alternatively, reply **B** to paste a new resume text for me to review."
                            ),
                            "pdf_url": None,
                            "docx_url": None
                        }
                    
                    # Build improvement suggestions based on profile data
                    score = ats_score
                    skill_list = [s.strip() for s in profile_skills.split(",") if s.strip()]
                    
                    # Score-based feedback
                    if score >= 80:
                        score_feedback = f"🟢 **ATS Score: {score}/100** — Excellent! Your resume is well-optimized."
                        score_tip = "Your resume is in great shape. Focus on tailoring it to specific job descriptions."
                    elif score >= 60:
                        score_feedback = f"🟡 **ATS Score: {score}/100** — Good, but there's room to improve."
                        score_tip = "Add more quantified achievements (e.g., 'Increased performance by 40%') and ensure section headers are clear."
                    elif score >= 40:
                        score_feedback = f"🟠 **ATS Score: {score}/100** — Needs Improvement."
                        score_tip = "Your resume is missing key sections or lacks impactful content. See tips below."
                    else:
                        score_feedback = f"🔴 **ATS Score: {score}/100** — Critical issues detected."
                        score_tip = "Your resume needs significant improvements in structure, contact info, and experience descriptions."
                    
                    # Skill feedback
                    if len(skill_list) >= 10:
                        skill_feedback = f"✅ **Skills:** {len(skill_list)} skills detected — great coverage!"
                    elif len(skill_list) >= 5:
                        skill_feedback = f"🟡 **Skills:** {len(skill_list)} skills detected. Add 5-10 more for a stronger profile."
                    else:
                        skill_feedback = f"⚠️ **Skills:** Only {len(skill_list)} skills detected. Add at least 10-15 technical skills."
                    
                    improvements = []
                    if score < 80:
                        improvements.append("📌 **Add quantified metrics** — e.g., 'Built app used by 5,000 users'")
                    if score < 60:
                        improvements.append("📌 **Strengthen Experience section** — use action verbs like Built, Led, Reduced, Deployed")
                    if score < 40:
                        improvements.append("📌 **Add Contact Info** — ensure email, phone, LinkedIn are present")
                        improvements.append("📌 **Add clear section headers**: Experience, Education, Skills, Projects")
                    if len(skill_list) < 8:
                        improvements.append("📌 **Expand Skills** — include frameworks, tools, and databases you know")
                    improvements.append("📌 **Add GitHub and project links** — boosts ATS score significantly")
                    improvements.append("📌 **Tailor your resume to each job** — match keywords from the job description")
                    
                    tips_text = "\n".join(improvements)
                    
                    return {
                        "reply": (
                            f"## 📊 Your Profile Resume Analysis\n\n"
                            f"{score_feedback}\n\n"
                            f"{skill_feedback}\n\n"
                            f"### 🔧 Improvement Tips:\n{tips_text}\n\n"
                            f"### 💡 Summary:\n{score_tip}\n\n"
                            f"---\n*Type **build** to create an improved resume, or **menu** to go back.*"
                        ),
                        "pdf_url": None,
                        "docx_url": None
                    }
                
                elif is_new:
                    session["_rate_step"] = 1  # new resume path
                    return {
                        "reply": (
                            "Please **paste your resume text** below and I'll analyze it thoroughly! 📋\n\n"
                            "You can paste the full text content of your resume (copy all text from your PDF or Word document)."
                        ),
                        "pdf_url": None,
                        "docx_url": None
                    }
                else:
                    return {
                        "reply": "Please reply with **A** for your profile resume, or **B** to paste a new resume.",
                        "pdf_url": None,
                        "docx_url": None
                    }
            
            # Rate step 1: user pasted a new resume text
            elif rate_step == 1:
                resume_text = user_messages[-1].content.strip() if user_messages else ""
                if len(resume_text.split()) < 30:
                    return {
                        "reply": "⚠️ The resume text seems too short. Please paste the full text content of your resume.",
                        "pdf_url": None,
                        "docx_url": None
                    }
                
                session["_rate_step"] = 2
                
                # Import scoring functions from main app context
                try:
                    import sys, os
                    sys.path.insert(0, os.path.dirname(__file__))
                    from main import calculate_ats_score, extract_skills, generate_suggestions, get_missing_skills
                    
                    found_skills = extract_skills(resume_text)
                    ats_result = calculate_ats_score(resume_text, found_skills)
                    score = ats_result["score"]
                    suggestions = generate_suggestions(ats_result, found_skills, get_missing_skills(found_skills, "software developer"))
                    
                    suggestion_text = "\n".join([f"• {s}" for s in suggestions[:8]])
                    skill_text = ", ".join(found_skills[:12]) if found_skills else "None detected"
                    
                    return {
                        "reply": (
                            f"## 📊 Resume Analysis Results\n\n"
                            f"**ATS Score: {score}/100**\n\n"
                            f"**Detected Skills:** {skill_text}\n\n"
                            f"### 🔧 Improvement Suggestions:\n{suggestion_text}\n\n"
                            f"---\n*Type **build** to create an optimized resume, or **menu** to return.*"
                        ),
                        "pdf_url": None,
                        "docx_url": None
                    }
                except Exception as e:
                    # Fallback to generic feedback
                    wc = len(resume_text.split())
                    return {
                        "reply": (
                            f"## 📊 Resume Analysis\n\n"
                            f"**Word Count:** {wc} words {'✅ (good length)' if 300 <= wc <= 1200 else '⚠️ (aim for 400-800 words)'}\n\n"
                            f"### Key Recommendations:\n"
                            f"• Ensure email, phone, and LinkedIn are present\n"
                            f"• Add strong action verbs (Led, Built, Developed, Deployed)\n"
                            f"• Include project links and GitHub\n"
                            f"• Quantify your impact (e.g., '10K users', '40% faster')\n\n"
                            f"---\n*Type **build** to create a polished resume, or **menu** to return.*"
                        ),
                        "pdf_url": None,
                        "docx_url": None
                    }
        
        # ══ BUILD MODE ═══════════════════════════════════════════════════════════
        if mode == "build":
            step = session.get("_step", 0)
            
            # Check if user wants to restart
            if last_msg in ["restart", "start over", "new", "reset"]:
                session["_step"] = 0
                return {
                    "reply": "Starting fresh! 🔄\n\n" + STEPS[0]["question"],
                    "pdf_url": None,
                    "docx_url": None
                }
            
            if step < len(STEPS):
                current_key = STEPS[step]["key"]
                _store_answer(uid, current_key, user_messages[-1].content.strip() if user_messages else "")
                
                # Auto-fill skills from profile
                if current_key == "skills" and len(user_messages[-1].content.split(",")) < 3:
                    profile_skills = session.get("_profile_skills", "")
                    if profile_skills:
                        orig = user_messages[-1].content.strip()
                        session["skills"] = (orig + ", " + profile_skills) if orig.lower() != "skip" else profile_skills

                new_step = session.get("_step", 0)
                if new_step < len(STEPS):
                    progress = f"[{new_step}/{len(STEPS)}]"
                    return {
                        "reply": f"Got it! {progress}\n\n{STEPS[new_step]['question']}",
                        "pdf_url": None,
                        "docx_url": None
                    }
                else:
                    # Generate resume
                    resume_data = _build_resume_data(uid)
                    import random, uuid as _uuid
                    template = random.choice(RESUME_TEMPLATES)
                    file_id = str(_uuid.uuid4())
                    
                    try:
                        pdf_url = generate_pdf(resume_data, file_id, template)
                        docx_url = generate_docx(resume_data, file_id)
                        _session_data[uid] = {"_step": 0, "_mode": None}
                        return {
                            "reply": (
                                f"🎉 Your resume is ready using the **{template['name']}** template!\n\n"
                                f"Download it below. Type **restart** to create another."
                            ),
                            "pdf_url": pdf_url,
                            "docx_url": docx_url
                        }
                    except Exception as e:
                        import traceback
                        traceback.print_exc()
                        return {
                            "reply": f"Error generating resume: {str(e)}. Please try again.",
                            "pdf_url": None,
                            "docx_url": None
                        }
            else:
                _session_data[uid] = {"_step": 0, "_mode": None}
                return {
                    "reply": "Your resume has already been generated! Type **restart** to create a new one, or **menu** to go back.",
                    "pdf_url": None,
                    "docx_url": None
                }
        
        # Fallback
        return {
            "reply": "I'm not sure what you'd like to do. Type **menu** to see options.",
            "pdf_url": None,
            "docx_url": None
        }
    
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))



